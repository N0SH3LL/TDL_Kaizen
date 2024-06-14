import SCCREAD
import SCCCHECK
import SCCTABLES
import SPLITBPER
import FILEGRAB
import argparse
import os
import json
import shutil
import sys
import re
import docx
from datetime import datetime

def read_json(filename): #needed for opening progress.json
    with open(filename, 'r') as file:
        return json.load(file)

def write_not_gathered_file(output_filename='Not_gathered.txt'): # obsolete
    progress_data = read_json('progress.json')
    all_bper_dict = progress_data.get('BPERs', {})
    all_doc_dict = progress_data.get('Documents', {})
    all_attestation_dict = progress_data.get('Attestations', {})
    
    not_gathered_info = {}

    for data_dict in [all_attestation_dict, all_bper_dict, all_doc_dict]:
        for key, entries in data_dict.items():
            for value in entries:
                # Check for 'Gathered' being blank or 'TLA' being True in all_bper_dict
                if value.get('Gathered') == '' or (value.get('TLA') == True and data_dict is all_bper_dict):
                    scc_name = value['SCC']
                    if scc_name not in not_gathered_info:
                        not_gathered_info[scc_name] = []
                    detail = value.get('Attestation num', '') or value.get('BPER name', '') or value.get('Doc name', '')
                    if detail:
                        not_gathered_info[scc_name].append(detail)

    # Write to the output file
    with open(output_filename, 'w') as output_file:
        for scc_name, details in sorted(not_gathered_info.items()):
            output_file.write(f"{scc_name}\n")
            for detail in sorted(details):
                output_file.write(f"\t{detail}\n")
            output_file.write("\n")

def update_dict(all_dict, new_entries): # check the dictionary to make sure any value being added is unique
    for key, value in new_entries.items():
        if key not in all_dict:
            all_dict[key] = []
        all_dict[key].append(value)

def replace_text_in_docx(doc_path, replacements): #supports filling in some of the templates
    doc = docx.Document(doc_path)

    # Replace in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for key, val in replacements.items():
                            if key in run.text:
                                run.text = run.text.replace(key, val)

    doc.save(doc_path)

def create_directories(project_dir): # creates the directories in the project directory
    progress_file_path = os.path.join(project_dir, 'progress.json') # put the progress.json file for the project in the project directory root

    try:
        with open(progress_file_path, 'r') as file: # grab the settings from progress.json
            progress_data = json.load(file)
            scc_dict = progress_data.get('SCC', {})
            project_settings = progress_data.get('Program Settings', {})
            project_directory = project_settings.get('Project Directory', project_dir)
    except (IOError, json.JSONDecodeError) as e:
        print(f"Failed to read or parse the progress.json file: {e}")
        return

    subdirs = ["Attestations", "Automated", "Exceptions and Deviations", "Manual", "Supporting Documents"] # Creates these subdirectories

    for scc_path, details in scc_dict.items():
        scc_name = details['SCC']
        sanitized_scc_name = re.sub(r'_\d{2}$', '', scc_name).strip()  # remove trailing "_**"
        main_dir_path = os.path.join(project_directory, sanitized_scc_name)

        if not details.get('Directory built', False):  # Check if directory is not already built
            os.makedirs(main_dir_path, exist_ok=True)

            for subdir in subdirs: # make subdirs above
                os.makedirs(os.path.join(main_dir_path, subdir), exist_ok=True)

            print(f"Created directory structure for: {sanitized_scc_name}")
            details['Directory built'] = True  # Mark as built

    # Save the updated progress data to progress.json
    with open(progress_file_path, 'w') as file:
        json.dump(progress_data, file, indent=4)

    print("All directories are created.")

def build_templates(method_dict, project_dir, template_dir): #Creates the directories, loads and modifies templates
    templates_to_copy = {
        "Manual": ["Teamname-Manual_Control_Evidence.xlsx"],
        "Attestations": ["Teamname-Attestation_Evidence.xlsx"],
        "Supporting Documents": ["Teamname-Document_Evidence.xlsx"],
        "Automated": ["Teamname-EvidenceValidation.xlsx"],
        "Root": ["Teamname-DeviceGapList.xlsx", "Teamname-Remediation.xlsx"]
    }

    directories_to_process = set()
    for stig_id, details in method_dict.items():
        scc_name = details['SCC']
        sanitized_scc_name = re.sub(r'_\d{2}$', '', scc_name).strip()
        directories_to_process.add(sanitized_scc_name)

    for sanitized_scc_name in directories_to_process:
        main_dir_path = os.path.join(project_dir, sanitized_scc_name)
        os.makedirs(main_dir_path, exist_ok=True)

        subdirs = ["Attestations", "Automated", "Exceptions and Deviations", "Manual", "Supporting Documents"]
        subdirs_paths = {subdir: os.path.join(main_dir_path, subdir) for subdir in subdirs}
        for subdir, subdir_path in subdirs_paths.items():
            os.makedirs(subdir_path, exist_ok=True)

            if subdir == "Manual":
                automated_info_path = os.path.join(subdir_path, 'Automated Info')
                os.makedirs(automated_info_path, exist_ok=True)

                # Create a copy of "Manual Screenshot Template.docx" for each manual STIG ID
                for stig_id, details in method_dict.items():
                    if details['Evidence method'].lower() == 'manual':
                        sanitized_stig_id = re.sub(r'[<>:"/\\|?*]', '', stig_id).strip()
                        template_path = os.path.join(template_dir, "Manual Screenshot Template.docx")
                        if os.path.isfile(template_path):
                            stig_template_filename = f"{sanitized_stig_id}.docx"
                            stig_template_path = os.path.join(subdir_path, stig_template_filename)
                            try:
                                shutil.copy2(template_path, stig_template_path)
                                print(f"Copied Manual Screenshot Template to {stig_template_path}")

                                replacements = {
                                    "FILENAMEINSERT": sanitized_scc_name,
                                    "STIGIDINSERT": stig_id
                                }
                                replace_text_in_docx(stig_template_path, replacements)
                            
                            except Exception as e:
                                print(f"Error copying manual template for STIG ID {stig_id}: {e}")

        # Copy and rename other templates
        for subdir, templates in templates_to_copy.items():
            for template in templates:
                template_path = os.path.join(template_dir, template)
                if os.path.isfile(template_path):
                    # If 'Root', the destination is the main directory; otherwise, it's a subdirectory
                    if subdir == "Root":
                        dest_path = os.path.join(project_dir, sanitized_scc_name, template.replace("Teamname", sanitized_scc_name))
                    else:
                        dest_path = os.path.join(subdirs_paths[subdir], template.replace("Teamname", sanitized_scc_name))

                    try:
                        shutil.copy2(template_path, dest_path)
                        print(f"Copied {template} to {dest_path}")
                    except Exception as e:
                        print(f"Error copying {template}: {e}")
                else:
                    print(f"Template file not found: {template_path}")

        print(f"Created directory for: {sanitized_scc_name}")

def convert_datetime_to_string(obj): # converts datetime values to string because it pitches a fit if not done
    if isinstance(obj, dict):
        return {key: convert_datetime_to_string(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [convert_datetime_to_string(item) for item in obj]
    elif isinstance(obj, datetime):
        return obj.isoformat()
    return obj

def save_master_dicts(data, filename): # save the dictionaries after converting the values to string
    converted_data = convert_datetime_to_string(data)
    with open(filename, 'w') as file:
        json.dump(converted_data, file, indent=4)

def move_sccs_to_folders(project_dir, file_extensions): # moves the scc's into their respetive directories ****NOT USED IN GUI, NEEDS TO COPY INSTEAD****
    for file in os.listdir(project_dir):
        file_path = os.path.join(project_dir, file)
        file_name, file_ext = os.path.splitext(file)

        # Apply the sanitization only for Excel files
        if file_ext in ['.xlsx', '.xls']:
            sanitized_folder_name = re.sub(r'_\d{2}$', '', file_name).strip()
        else:
            sanitized_folder_name = file_name

        if file_ext in file_extensions:
            dest_folder_path = os.path.join(project_dir, sanitized_folder_name)

            if os.path.exists(dest_folder_path) and os.path.isdir(dest_folder_path):
                shutil.move(file_path, os.path.join(dest_folder_path, file))
            else:
                print(f"Destination folder not found for {file}, expected at {dest_folder_path}")

def build_progress_json(directory_path, project_dir):
    #Master dictionaries
    all_bper_dict = {}
    all_doc_dict = {}
    all_attestation_dict = {}
    scc_data_dict = {}
    checks_data_dict = {}

    #Loop for all excel files in directory
    for file in os.listdir(directory_path):
        if file.endswith('.xlsx') or file.endswith('.xls'):
            file_path = os.path.join(directory_path, file)

            #This creates the dictionaries using SCCREAD (bper, docs, atts, controls)
            bper_dict, doc_dict, attestation_dict, method_dict = SCCREAD.process_excel_file(file_path)

            #This creates the scc_info dictionary (first page check, etc)
            scc_info = SCCCHECK.process_scc_file(file_path)

            #stores in master dict
            update_dict(all_bper_dict, bper_dict)
            update_dict(all_attestation_dict, attestation_dict)
            update_dict(all_doc_dict, doc_dict)

            # Update scc_data_dict and checks_data_dict
            scc_data_dict[file_path] = scc_info
            for stig_id, details in method_dict.items():
                # Extract SCC name from the file path and remove extension and trailing "_**"
                scc_name = os.path.splitext(os.path.basename(file_path))[0]
                scc_name = re.sub(r'_\d{2}$', '', scc_name).strip()
                checks_data_dict[stig_id] = {
                    'SCC': scc_name,
                    'Evidence method': details['Evidence Method']
                }

    # Save progress to progress.json
    progress_data = {
        'BPERs': all_bper_dict,
        'Attestations': all_attestation_dict,
        'Documents': all_doc_dict,
        'SCC': scc_data_dict,
        'Checks': checks_data_dict,
        'Program Settings': {
            'Project Directory': project_dir,
            'Directories Built': False,
            'Templates Built': False,
            'Gather and Sort Date': '',
            'Doc Tracker Update': '',
            'Pull Info Date': '',
            'Checklists generated':''
        }
    }
    with open(os.path.join(project_dir, 'progress.json'), 'w') as file:
        json.dump(convert_datetime_to_string(progress_data), file, indent=4)

def update_bper_dict(master_directory):
    progress_data = read_json('progress.json')
    bper_dict = progress_data.get('BPERs', {})

    for key, value_list in bper_dict.items():
        for value in value_list:
            scc_name = value['SCC']
            file_path = os.path.join(master_directory, scc_name, "Exceptions and Deviations", f"{key}.pdf")
            if os.path.isfile(file_path):
                valid_to_date, approval_status, tla_present = FILEGRAB.extract_BPER_info(file_path)
                value['Valid to'] = valid_to_date
                value['Approval Status'] = approval_status
                value['TLA'] = tla_present
                print(f"Updated BPER '{key}' - 'Valid to': {valid_to_date}, 'Approval Status': {approval_status}, 'TLA': {tla_present}")
            else:
                print(f"File not found for BPER: {key}")

    progress_data['BPERs'] = bper_dict
    with open('progress.json', 'w') as file:
        json.dump(convert_datetime_to_string(progress_data), file, indent=4)

def update_attestation_dict(master_directory):
    progress_data = read_json('progress.json')
    attestation_dict = progress_data.get('Attestations', {})

    for key, value_list in attestation_dict.items():
        for value in value_list:
            scc_name = value['SCC']
            file_path = os.path.join(master_directory, scc_name, "Attestations", f"{key}.pdf")
            if os.path.isfile(file_path):
                approval_status, approval_date = FILEGRAB.extract_attest_info(file_path)
                value['Approval Status'] = approval_status
                value['Valid to'] = approval_date
                print(f"Updated '{key}' with status: {approval_status} and date: {approval_date}")
            else:
                print(f"File not found for Attestation: {key}")

    progress_data['Attestations'] = attestation_dict
    with open('progress.json', 'w') as file:
        json.dump(convert_datetime_to_string(progress_data), file, indent=4)

def update_doc_dict(master_directory):
    progress_data = read_json('progress.json')
    doc_dict = progress_data.get('Documents', {})

    for key, value_list in doc_dict.items():
        for value in value_list:
            scc_name = value['SCC']
            file_path = os.path.join(master_directory, scc_name, "Supporting Documents", f"{key}.docx")
            if os.path.isfile(file_path):
                most_recent_date = FILEGRAB.extract_Doc_info(file_path)
                if most_recent_date:
                    value['Last update'] = most_recent_date
                    print(f"Updated '{key}' with 'Last update': {most_recent_date}")
            else:
                print(f"File not found for Document: {key}")

    progress_data['Documents'] = doc_dict
    with open('progress.json', 'w') as file:
        json.dump(convert_datetime_to_string(progress_data), file, indent=4)

def main():
    parser = argparse.ArgumentParser(description='Main script to process SCC files.')
    parser.add_argument('directory_path', type=str, help='Path to the directory containing Excel files')
    parser.add_argument('--progress', action='store_true', help='Load progress from progress.json')
    args = parser.parse_args()

    if not os.path.isdir(args.directory_path):
        print(f"Directory not found: {args.directory_path}")
        return

    #Checks and splits bulk BPERs
    All_BPER_path = os.path.join(args.directory_path, 'All BPERs')
    SPLITBPER.process_directory(All_BPER_path)
    master_directory = args.directory_path
    
    # Load progress from progress.json if --progress flag is set
    if args.progress:
        progress_file = 'progress.json'
        if os.path.exists(progress_file):
            with open(progress_file, 'r') as file:
                progress_data = json.load(file)
        else:
            print(f"Progress file {progress_file} not found. Starting from scratch.")
            progress_data = {
                'BPERs': {},
                'Attestations': {},
                'Documents': {},
                'SCC': {},
                'Checks': {},
                'Program Settings': {
                    'Project Directory': master_directory,
                    'Directories Built': False,
                    'Templates Built': False,
                    'Gather and Sort Date': '',
                    'Doc Tracker Update': ''
                }
            }
    else:
        progress_data = {
            'BPERs': {},
            'Attestations': {},
            'Documents': {},
            'SCC': {},
            'Checks': {},
            'Program Settings': {
                'Project Directory': master_directory,
                'Directories Built': False,
                'Templates Built': False,
                'Gather and Sort Date': '',
                'Doc Tracker Update': ''
            }
        }

    #Loop for all excel files in directory
    for file in os.listdir(args.directory_path):
        if file.endswith('.xlsx') or file.endswith('.xls'):
            file_path = os.path.join(args.directory_path, file)

            # Check if the file has already been processed
            if file_path in progress_data['SCC']:
                print(f"Skipping already processed file: {file_path}")
                continue

            #This creates the dictionaries using SCCREAD (bper, docs, atts, controls)
            bper_dict, doc_dict, attestation_dict, method_dict = SCCREAD.process_excel_file(file_path)
            #This creates the scc_info dictionary (first page check, etc)
            scc_info = SCCCHECK.process_scc_file(file_path)

            #This creates the directories for each SCC
            build_templates(master_directory)

            #This grabs the files and updates the dictionaries
            base_directories = {'bper': 'All BPERs', 'doc': 'All Docs', 'attestation': 'All Attestations'}
            FILEGRAB.update_dictionaries_and_copy_files(bper_dict, doc_dict, attestation_dict, base_directories, master_directory)

            scc_name_without_extension = os.path.splitext(os.path.basename(file_path))[0]
            scc_name_without_extension = re.sub(r'_\d{2}$', '', scc_name_without_extension).strip()
            print(f"\nGathered documents for {scc_name_without_extension}")

            #This writes the checklist for each SCC
            SCCTABLES.write_checklist(bper_dict, doc_dict, attestation_dict, method_dict, scc_info, master_directory)

            #stores in master dict
            update_dict(progress_data['BPERs'], bper_dict)
            update_dict(progress_data['Attestations'], attestation_dict)
            update_dict(progress_data['Documents'], doc_dict)

            # Update scc_data_dict and checks_data_dict
            progress_data['SCC'][file_path] = scc_info
            for stig_id, details in method_dict.items():
                # Extract SCC name from the file path and remove extension and trailing "_**"
                scc_name = os.path.splitext(os.path.basename(file_path))[0]
                scc_name = re.sub(r'_\d{2}$', '', scc_name).strip()
                progress_data['Checks'][stig_id] = {
                    'SCC': scc_name,
                    'Evidence method': details['Method']
                }

    #move text and SCC's into their folders
    file_extensions = ['.xlsx', '.txt']
    move_sccs_to_folders(master_directory, file_extensions)

    #saves master dicts
    save_master_dicts(progress_data['Attestations'], 'all_attestation_dict.json')
    save_master_dicts(progress_data['BPERs'], 'all_bper_dict.json')
    save_master_dicts(progress_data['Documents'], 'all_doc_dict.json')

    #build not gathered list
    write_not_gathered_file()

    # Save progress to progress.json
    with open('progress.json', 'w') as file:
        json.dump(convert_datetime_to_string(progress_data), file, indent=4)


if __name__ == "__main__":
    main()
    print('done')
